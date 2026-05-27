from __future__ import annotations

from datetime import datetime
from pathlib import Path
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\hfk47\Documents\noderasoftware")
OUT = ROOT / "HotelOps_Komut_Egitim_Dokumani.docx"


ACCENT = "1F4E79"
LIGHT = "EAF2F8"
GRAY = "F4F6F8"
DARK = "1F2933"
MUTED = "64748B"
BORDER = "CBD5E1"


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="4"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for i, width in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = width
                row.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                set_cell_border(row.cells[i])


def paragraph_border_bottom(paragraph, color="D7DEE8", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Sayfa ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 116, 139)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r = paragraph.add_run()
    r._r.append(fld_begin)
    r._r.append(instr_text)
    r._r.append(fld_end)


def add_code_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.style = "CodeBlock"
    wrapped_lines: list[str] = []
    for raw in text.splitlines() or [""]:
        if len(raw) <= 110:
            wrapped_lines.append(raw)
        else:
            wrapped_lines.extend(textwrap.wrap(raw, width=110, subsequent_indent="    ", break_long_words=False, break_on_hyphens=False) or [raw])
    for line_no, line in enumerate(wrapped_lines or [""]):
        if line_no:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(31, 41, 55)
    return p


def add_callout(doc: Document, title: str, body: str, fill="FFF7ED"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.4)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell, color="FED7AA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(154, 52, 18)
    run.font.size = Pt(10)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    p2.runs[0].font.size = Pt(10)


def add_command_table(doc: Document, rows: list[tuple[str, str, str]], title: str | None = None):
    if title:
        heading = doc.add_heading(title, level=3)
        heading.paragraph_format.keep_with_next = True
    intro = doc.add_paragraph("Aşağıdaki komutlar aynı başlık altında kullanılan pratik komut setidir.")
    intro.paragraph_format.keep_with_next = True
    for row in rows:
        if len(row) == 2:
            command, purpose = row
            note = ""
        else:
            command, purpose, note = row
        label = doc.add_paragraph()
        label.paragraph_format.space_after = Pt(1)
        label.paragraph_format.keep_with_next = True
        r = label.add_run("Komut")
        r.bold = True
        r.font.color.rgb = RGBColor(31, 78, 121)
        add_code_paragraph(doc, command)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("Ne işe yarar: ")
        run.bold = True
        p.add_run(purpose)
        if note:
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Pt(10)
            p2.paragraph_format.space_after = Pt(6)
            run2 = p2.add_run("Not: ")
            run2.bold = True
            p2.add_run(note)
    doc.add_paragraph()


def add_script_listing(doc: Document, path: Path, title: str, explanation: str):
    doc.add_heading(title, level=2)
    p = doc.add_paragraph(explanation)
    p.paragraph_format.space_after = Pt(6)
    add_code_paragraph(doc, path.read_text(encoding="utf-8"))


def setup_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(31, 41, 55)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(5)

    title = styles["Title"]
    title.font.name = "Arial"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 78, 121)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    if "CodeBlock" not in styles:
        code = styles.add_style("CodeBlock", 1)
    else:
        code = styles["CodeBlock"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor(15, 23, 42)
    code.paragraph_format.left_indent = Pt(8)
    code.paragraph_format.right_indent = Pt(8)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.0


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def build():
    doc = Document()
    setup_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header.paragraphs[0]
    header.text = "HotelOps Kurulum ve Komut Eğitim Dokümanı"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(100, 116, 139)
    paragraph_border_bottom(header)
    add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph(style="Title")
    title.add_run("HotelOps Kurulum ve Komut Eğitim Dokümanı")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = subtitle.add_run("IIS, Next.js, Node.js API, PostgreSQL, Prisma ve Windows Service adımlarının açıklamalı komut rehberi")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(71, 85, 105)

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    meta_rows = [
        ("Proje klasörü", r"C:\Users\hfk47\Documents\noderasoftware"),
        ("Canlı IIS klasörü", r"C:\inetpub\wwwroot"),
        ("Doküman tarihi", datetime.now().strftime("%d.%m.%Y %H:%M")),
        ("Hazırlanma amacı", "Yapay zeka olmadan aynı kurulumu ve kontrolleri elle yapabilmek"),
    ]
    for i, (k, v) in enumerate(meta_rows):
        meta.cell(i, 0).text = k
        meta.cell(i, 1).text = v
        shade_cell(meta.cell(i, 0), LIGHT)
        meta.cell(i, 0).paragraphs[0].runs[0].bold = True
    set_table_width(meta, [Inches(1.8), Inches(4.6)])

    add_callout(
        doc,
        "Güvenlik notu",
        "Bu doküman eğitim amaçlıdır. Parola, JWT secret, database şifresi ve dış IP gibi bilgiler canlı sistemlerde benzersiz ve gizli tutulmalıdır. Komutlarda gerekli yerlerde örnek değerler veya yerel proje değerleri gösterilmiştir.",
    )

    doc.add_heading("1. Büyük Resim", level=1)
    for item in [
        "IIS port 80 üzerinden frontend dosyalarını C:\\inetpub\\wwwroot klasöründen yayınlar.",
        "Frontend statik Next.js çıktısıdır; tarayıcıda çalışır ve aynı domain altındaki /api yoluna istek atar.",
        "IIS URL Rewrite + ARR, /api isteklerini arkadaki Node.js API'ye yollar.",
        "Node.js API 4000 portunda çalışır ve PostgreSQL veritabanına Prisma ile bağlanır.",
        "PostgreSQL servisi veriyi tutar; kullanıcılar, roller, iş emirleri, audit log ve oturum kayıtları buradadır.",
        "HotelOpsApi Windows servisi API'nin bilgisayar açılışında otomatik başlamasını sağlar.",
    ]:
        add_bullet(doc, item)

    add_command_table(
        doc,
        [
            ("http://127.0.0.1/", "IIS'in yayınladığı frontend ana sayfasını açar.", "Sunucunun kendi içinde hızlı kontrol için."),
            ("http://127.0.0.1/api/health", "IIS proxy üzerinden API ve DB sağlığını kontrol eder.", "Canlı kullanıcıların gördüğü yolun API'ye ulaşıp ulaşmadığını gösterir."),
            ("http://127.0.0.1:4000/health", "Node API'ye direkt gider.", "Proxy değil, doğrudan API çalışıyor mu anlamak için."),
            ("SEED_DEFAULT_PASSWORD", "İlk kurulum parolası ortam değişkeninden alınır.", "Canlı sistemde sabit demo parolası kullanılmaz."),
        ],
        "Hızlı Sağlık Kontrol Noktaları",
    )

    doc.add_heading("2. Komut Çalıştırma Kuralları", level=1)
    for item in [
        "PowerShell komutları proje klasöründe çalıştırıldı: C:\\Users\\hfk47\\Documents\\noderasoftware.",
        "IIS, servis kurulumu, Chocolatey kurulumu ve C:\\inetpub yazma işlemleri yönetici yetkisi ister.",
        "Node/Next build komutları normal kullanıcıyla çalışabilir.",
        "Canlı klasöre dosya basmadan önce her zaman yedek alınmalıdır.",
        "Dış dünyaya hata göstermemek için önce 127.0.0.1 üzerinde doğrulama yapılmalıdır.",
    ]:
        add_bullet(doc, item)

    add_command_table(
        doc,
        [
            ("Get-ChildItem -Path . -Force", "Klasördeki dosya ve dizinleri listeler.", "Klasörün doğru yer olup olmadığını görmek için."),
            ("rg --files", "Projede bulunan tüm dosyaları hızlı listeler.", "Kod tabanının şeklini görmek için grep yerine tercih edildi."),
            ("Get-Content -Path package.json", "Ana monorepo scriptlerini okur.", "Build, dev, api:dev gibi scriptleri görmek için."),
            ("Get-Content -Path apps\\web\\package.json", "Frontend script ve bağımlılıklarını okur.", "Next.js komutlarını ve port ayarlarını görmek için."),
            ("Get-Content -Path apps\\api\\package.json", "API script ve bağımlılıklarını okur.", "Express, Prisma, JWT, bcrypt gibi backend paketlerini görmek için."),
            ("Get-ChildItem -Path C:\\inetpub\\wwwroot -Force", "IIS canlı klasöründeki dosyaları listeler.", "Eski ASP.NET yayını mı yeni statik çıktı mı anlaşılır."),
            ("Get-Content -Path C:\\inetpub\\wwwroot\\Web.config", "IIS site konfigürasyonunu okur.", "Default document, rewrite/proxy kuralları ve MIME ayarları burada görülür."),
        ],
        "Keşif ve Dosya İnceleme Komutları",
    )

    doc.add_heading("3. Frontend Build ve Yayın", level=1)
    doc.add_paragraph(
        "Frontend Next.js ile geliştirildi. Son canlı durumda output: export kullanıldı; yani build sonrası apps\\web\\out klasörü IIS'in yayınlayabileceği statik HTML/CSS/JS dosyalarına dönüştü."
    )
    add_command_table(
        doc,
        [
            ("npm.cmd run typecheck --workspace @hotel-ops/web", "Frontend TypeScript kontrolü yapar.", "Kod değişikliğinden sonra tip hatası var mı görmek için."),
            ("npm.cmd run lint --workspace @hotel-ops/web", "Frontend lint kontrolü yapar.", "Kod kalitesi ve kullanılmayan değişkenleri yakalamak için."),
            ("npm.cmd run build --workspace @hotel-ops/web", "Next.js production/static build üretir.", "Canlıya basılacak apps\\web\\out klasörü bununla güncellenir."),
            ("Invoke-WebRequest -UseBasicParsing -Uri \"http://127.0.0.1/\" | Select-Object -ExpandProperty StatusCode", "IIS ana sayfasının 200 dönüp dönmediğini kontrol eder.", "Yayın sonrası ilk HTTP kontrolü."),
            ("Invoke-WebRequest -UseBasicParsing -Uri \"http://127.0.0.1/dashboard\" | Select-Object -ExpandProperty StatusCode", "Dashboard route'unun IIS üzerinden açıldığını kontrol eder.", "Statik route üretildi mi anlamak için."),
        ],
        "Frontend Doğrulama ve Build Komutları",
    )

    doc.add_heading("4. IIS Canlı Yayın ve Yedekleme", level=1)
    doc.add_paragraph(
        "Canlı klasör doğrudan silinmedi. Önce C:\\inetpub\\wwwroot-backups altında zaman damgalı yedek alındı, sonra yeni apps\\web\\out çıktısı wwwroot içine kopyalandı."
    )
    add_command_table(
        doc,
        [
            ("Test-Path C:\\inetpub\\wwwroot\\_next", "Yeni Next.js statik asset klasörü var mı kontrol eder.", "_next varsa yeni frontend yayındadır."),
            ("Test-Path C:\\inetpub\\wwwroot\\Dashboard.aspx", "Eski ASP.NET dosyaları duruyor mu kontrol eder.", "Dashboard.aspx varsa eski ASP.NET yayını hâlâ klasörde olabilir."),
            ("Get-ChildItem -Path C:\\inetpub\\wwwroot-backups -Force | Sort-Object LastWriteTime -Descending | Select-Object -First 5 Name,LastWriteTime", "Son canlı yedekleri listeler.", "Rollback gerekirse hangi yedeğin kullanılacağını görmek için."),
            ("powershell.exe -NoProfile -ExecutionPolicy Bypass -File \"scripts\\deploy-iis-static.ps1\"", "Yeni statik frontend çıktısını IIS wwwroot'a kopyalar.", "Yönetici olarak çalıştırılmalıdır; script önce yedek alır."),
        ],
        "IIS Yayın Komutları",
    )

    doc.add_heading("5. IIS Reverse Proxy Kurulumu", level=1)
    doc.add_paragraph(
        "İlk başta frontend API'ye :4000 portundan gitmeye çalışıyordu. Dış kullanıcıda bu port kapalı olabildiği için IIS'e URL Rewrite + Application Request Routing kuruldu. Artık frontend /api yolunu kullanıyor; IIS bunu arkada 127.0.0.1:4000 adresine yönlendiriyor."
    )
    add_command_table(
        doc,
        [
            ("choco search urlrewrite --limit-output", "Chocolatey'de URL Rewrite paketini arar.", "Paket adı UrlRewrite olarak bulundu."),
            ("choco search application-request-routing --limit-output", "ARR paketini arar.", "Paket adı iis-arr olarak bulundu."),
            ("choco install UrlRewrite iis-arr --yes --no-progress", "IIS URL Rewrite ve ARR modüllerini kurar.", "Yönetici olarak çalıştırılır."),
            ("& \"$env:windir\\system32\\inetsrv\\appcmd.exe\" set config -section:system.webServer/proxy /enabled:\"True\" /preserveHostHeader:\"True\" /commit:apphost", "ARR proxy özelliğini IIS genelinde açar.", "/api yönlendirmesinin çalışması için gerekir."),
            ("iisreset", "IIS'i yeniden başlatır.", "Yeni modüller ve proxy ayarı devreye girsin diye."),
            ("Invoke-WebRequest -UseBasicParsing -Uri \"http://127.0.0.1/api/health\"", "IIS üzerinden API proxy sağlığını test eder.", "502 dönüyorsa IIS API'ye ulaşamıyor demektir."),
        ],
        "IIS Proxy Komutları",
    )

    add_code_paragraph(
        doc,
        """Web.config içindeki kritik proxy kuralı:
<rewrite>
  <rules>
    <rule name="HotelOps API Proxy" stopProcessing="true">
      <match url="^api/(.*)" />
      <action type="Rewrite" url="http://127.0.0.1:4000/{R:1}" />
    </rule>
  </rules>
</rewrite>""",
    )

    doc.add_heading("6. PostgreSQL Kurulumu ve Veritabanı", level=1)
    add_command_table(
        doc,
        [
            ("where.exe psql", "psql komut satırı aracı PATH içinde var mı bakar.", "Kurulumdan önce yoktu; sonra PostgreSQL bin klasöründe bulundu."),
            ("Get-Service | Where-Object { $_.Name -like '*postgres*' -or $_.DisplayName -like '*Postgre*' }", "PostgreSQL servisi var mı kontrol eder.", "Servis yoksa kurulum/init gerekir."),
            ("choco install postgresql --yes --no-progress --params \"'/Password:<DB_PASSWORD> /Port:5432'\"", "PostgreSQL'i kurar.", "Eğitimde parola örnek olarak maskelendi; canlıda güçlü parola kullanın."),
            ("& 'C:\\Program Files\\PostgreSQL\\18\\bin\\initdb.exe' -D 'C:\\Program Files\\PostgreSQL\\18\\data' -U postgres --pwfile <pwfile> -E UTF8 --locale=C", "PostgreSQL data directory oluşturur.", "Kurulumda data klasörü boş kaldığı için gerekti."),
            ("Start-Service -Name postgresql-x64-18", "PostgreSQL servisini başlatır.", "Restart sonrası çalışmıyorsa ilk bakılacak komutlardan."),
            ("Get-Service -Name postgresql-x64-18 | Select-Object Name,Status,DisplayName", "PostgreSQL servis durumunu gösterir.", "Status Running olmalıdır."),
            ("$env:PGPASSWORD='<DB_PASSWORD>'; psql -h 127.0.0.1 -p 5432 -U postgres -c 'select version();'", "PostgreSQL'e bağlanıp sürüm testi yapar.", "DB dinliyor mu anlamak için."),
        ],
        "PostgreSQL Komutları",
    )

    add_callout(
        doc,
        "PostgreSQL hata notu",
        "Kurulumdan sonra servis 1067 ile düşmüştü. Event log'da postgresql.conf yok görünüyordu. Sebep data directory'nin initialize edilmemesiydi. initialize-postgres.ps1 bu işi düzeltmek için yazıldı.",
        fill="FEF2F2",
    )

    doc.add_heading("7. Prisma Şema, Migration ve Seed", level=1)
    add_command_table(
        doc,
        [
            ("npx.cmd prisma validate --schema prisma/schema.prisma", "Prisma şemasının geçerli olup olmadığını kontrol eder.", "DATABASE_URL yoksa burada hata verir."),
            ("npx.cmd prisma generate --schema prisma/schema.prisma", "Prisma Client üretir.", "Schema değişince API'nin yeni modelleri görmesi için."),
            ("npx.cmd prisma db push --schema prisma/schema.prisma", "Prisma şemasını PostgreSQL'e uygular.", "Geliştirme/ilk kurulum için hızlı senkronizasyon."),
            ("npm.cmd run seed --workspace @hotel-ops/api", "Seed verilerini DB'ye yükler.", "Demo kullanıcılar, roller, departmanlar, iş emirleri, takvim kayıtları."),
            ("$env:PGPASSWORD='<DB_PASSWORD>'; psql -h 127.0.0.1 -p 5432 -U hotelops -d hotelops", "hotelops kullanıcısıyla DB'ye bağlanır.", "Manuel SQL kontrolü yapmak için."),
        ],
        "Prisma ve Seed Komutları",
    )

    doc.add_heading("8. Node.js API Build, Test ve Çalıştırma", level=1)
    add_command_table(
        doc,
        [
            ("npm.cmd run build --workspace @hotel-ops/api", "TypeScript API kodunu dist/server.js olarak derler.", "Servise alınacak gerçek dosya budur."),
            ("node apps\\api\\dist\\server.js", "API'yi foreground'da çalıştırır.", "Servis sorunu varsa hatayı canlı görmek için."),
            ("Invoke-WebRequest -UseBasicParsing -Uri \"http://127.0.0.1:4000/health\"", "API'nin doğrudan ayakta olup olmadığını test eder.", "Burada hata varsa IIS değil API sorunudur."),
            ("Invoke-WebRequest -UseBasicParsing -Uri \"http://127.0.0.1/api/health\"", "IIS proxy üzerinden API'yi test eder.", "Burada 502 varsa API kapalı veya proxy bozuktur."),
            ("Get-Content -Path api-service.err -Tail 100", "API servis hata logunu okur.", "Path, DB, port veya Node hatalarını görmek için."),
            ("Get-Content -Path api-service.log -Tail 100", "API servis normal logunu okur.", "Dinleme başladı mı görmek için."),
        ],
        "API Komutları",
    )

    add_code_paragraph(
        doc,
        """Login ve bootstrap testi:
$body = @{ username='admin'; password='<kurulum-parolasi>' } | ConvertTo-Json
$login = Invoke-RestMethod -Uri 'http://127.0.0.1/api/auth/login' -Method Post -Body $body -ContentType 'application/json'
$bootstrap = Invoke-RestMethod -Uri 'http://127.0.0.1/api/bootstrap' -Headers @{ Authorization = "Bearer $($login.token)" }
[pscustomobject]@{ User=$bootstrap.user.username; Jobs=$bootstrap.jobs.Count; Users=$bootstrap.users.Count } | ConvertTo-Json""",
    )

    doc.add_heading("9. Windows Service Kurulumu", level=1)
    doc.add_paragraph(
        "Scheduled Task ile başlatma denendi; fakat görev Last Result: 1 ile düşüyordu. Kalıcı çözüm olarak NSSM kullanıldı ve HotelOpsApi gerçek Windows servisi yapıldı."
    )
    add_command_table(
        doc,
        [
            ("choco install nssm --yes --no-progress", "NSSM servis yöneticisini kurar.", "Node uygulamasını Windows Service yapmak için."),
            ("Get-Service -Name HotelOpsApi,postgresql-x64-18 | Select-Object Name,Status,DisplayName", "API ve PostgreSQL servis durumunu gösterir.", "İkisi de Running olmalıdır."),
            ("Start-Service -Name HotelOpsApi", "API servisini başlatır.", "Servis durmuşsa."),
            ("Restart-Service -Name HotelOpsApi", "API servisini yeniden başlatır.", "Kod değiştiyse veya sorun giderildiyse."),
            ("Stop-Service -Name HotelOpsApi", "API servisini durdurur.", "Bakım veya debug için."),
            ("& 'C:\\ProgramData\\chocolatey\\bin\\nssm.exe' get HotelOpsApi AppParameters", "Servisin hangi server.js dosyasını çalıştırdığını gösterir.", "Boşluklu path hatasını kontrol etmek için."),
            ("Get-NetTCPConnection -LocalPort 4000 -State Listen", "4000 portunda dinleyen süreç var mı gösterir.", "API gerçekten portu açmış mı anlamak için."),
        ],
        "Windows Service Komutları",
    )

    add_callout(
        doc,
        "Öğrenilen kritik ders",
        "Eski boşluklu proje yolu NSSM ilk kurulumda Node'a yanlış parametre geçmesine neden olmuştu. Proje artık boşluksuz C:\\Users\\hfk47\\Documents\\noderasoftware yolunda tutuluyor.",
        fill="EEF2FF",
    )

    doc.add_heading("10. Sorun Giderme Senaryoları", level=1)
    scenarios = [
        (
            "Login ekranı 'API servisine ulaşılamıyor' hatası veriyor",
            [
                ("Invoke-WebRequest -UseBasicParsing -Uri \"http://127.0.0.1/api/health\"", "Önce canlı yolun API'ye ulaşıp ulaşmadığını test et."),
                ("Invoke-WebRequest -UseBasicParsing -Uri \"http://127.0.0.1:4000/health\"", "Direkt API çalışıyor mu ayır."),
                ("Get-Service -Name HotelOpsApi,postgresql-x64-18", "Servislerin Running olup olmadığını kontrol et."),
                ("Restart-Service -Name HotelOpsApi", "API servisini yeniden başlat."),
            ],
        ),
        (
            "IIS /api/health 502 Bad Gateway dönüyor",
            [
                ("Invoke-WebRequest -UseBasicParsing -Uri \"http://127.0.0.1:4000/health\"", "API kapalıysa 502'nin sebebi budur."),
                ("Get-Content C:\\inetpub\\wwwroot\\Web.config", "Rewrite kuralı duruyor mu kontrol et."),
                ("iisreset", "IIS proxy modüllerini yeniden yükle."),
            ],
        ),
        (
            "PostgreSQL çalışmıyor",
            [
                ("Get-Service -Name postgresql-x64-18", "Servis durumunu kontrol et."),
                ("Start-Service -Name postgresql-x64-18", "Servisi başlat."),
                ("Get-Content postgres-diagnose.log", "Tanı scripti loglarını oku."),
            ],
        ),
        (
            "Yeni kod canlıya yansımadı",
            [
                ("npm.cmd run build --workspace @hotel-ops/web", "Yeni statik frontend build al."),
                ("powershell.exe -NoProfile -ExecutionPolicy Bypass -File \"scripts\\deploy-iis-static.ps1\"", "Build çıktısını IIS'e yayınla."),
                ("Ctrl + F5", "Tarayıcı cache'ini kır."),
            ],
        ),
    ]
    for title, rows in scenarios:
        add_command_table(doc, rows, title)

    doc.add_heading("11. Günlük Operasyon Kısa Rehberi", level=1)
    add_command_table(
        doc,
        [
            ("Get-Service -Name HotelOpsApi,postgresql-x64-18", "API ve DB servislerini kontrol et.", "Günlük ilk kontrol."),
            ("Invoke-WebRequest -UseBasicParsing -Uri \"http://127.0.0.1/api/health\"", "Canlı API yolunu kontrol et.", "ok:true ve db:up beklenir."),
            ("Restart-Service -Name HotelOpsApi", "API'yi yeniden başlat.", "Hata sonrası en hızlı düzeltme."),
            ("Get-Content api-service.err -Tail 100", "Son API hatalarını oku.", "Servis düşerse ilk bakılacak log."),
            ("Get-ChildItem C:\\inetpub\\wwwroot-backups | Sort-Object LastWriteTime -Descending", "Yedekleri gör.", "Rollback için."),
        ],
        "Günlük Kullanılacak Komutlar",
    )

    doc.add_heading("12. Canlıya Alma Akışı: Baştan Sona", level=1)
    steps = [
        "Kod değişikliğini yap.",
        "Frontend için typecheck, lint ve build çalıştır.",
        "API değiştiyse API build çalıştır.",
        "Prisma schema değiştiyse prisma validate, generate ve db push çalıştır.",
        "Seed gerekiyorsa seed çalıştır.",
        "API servisini yeniden başlat.",
        "Frontend out klasörünü IIS wwwroot'a yedekli yayınla.",
        "/api/health ve /dashboard test et.",
        "admin kullanıcısı ve kurulumda belirlenen parola ile login test et.",
    ]
    for i, step in enumerate(steps, start=1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)

    doc.add_heading("13. Oluşturulan Scriptler", level=1)
    doc.add_paragraph(
        "Aşağıdaki scriptler tekrar kullanılabilir otomasyonlardır. Eğitim için tam içerikleri kondu; gerçek canlı sistemde parolalar ve proje yolu ortama göre değiştirilmelidir."
    )

    script_items = [
        ("scripts\\deploy-iis-static.ps1", "IIS'e Yedekli Statik Yayın", "apps\\web\\out içeriğini C:\\inetpub\\wwwroot'a basmadan önce mevcut siteyi C:\\inetpub\\wwwroot-backups altına yedekler."),
        ("scripts\\install-postgres.ps1", "PostgreSQL Kurulum Scripti", "Chocolatey ile PostgreSQL kurar. Parola örnek amaçlıdır; canlıda güçlü parola gerekir."),
        ("scripts\\initialize-postgres.ps1", "PostgreSQL Init ve DB Oluşturma", "data directory yoksa initdb çalıştırır, servisi başlatır, hotelops rolü ve veritabanını oluşturur."),
        ("scripts\\install-iis-proxy-modules.ps1", "IIS URL Rewrite + ARR Kurulumu", "IIS proxy modüllerini kurar, ARR proxy'yi etkinleştirir ve IIS'i resetler."),
        ("scripts\\install-api-service.ps1", "HotelOpsApi Windows Service Kurulumu", "NSSM ile API'yi otomatik başlayan Windows servisi olarak kurar. Boşluklu yol sorunu için kısa path kullanır."),
        ("scripts\\start-api.ps1", "Manuel API Başlatma", "Node API'yi production modda başlatır ve log dosyalarına yazar."),
        ("scripts\\start-postgres.ps1", "PostgreSQL Servis Başlatma/Tanı", "PostgreSQL servisi yoksa kayıtlamayı dener, çalışmıyorsa pg_ctl loguyla tanı üretir."),
        ("scripts\\diagnose-postgres.ps1", "PostgreSQL Tanı Scripti", "Servis, data klasörü ve Windows event log bilgilerini postgres-diagnose.log dosyasına yazar."),
    ]
    for rel, title, expl in script_items:
        add_script_listing(doc, ROOT / rel, title, expl)

    doc.add_heading("14. Rollback Notları", level=1)
    add_command_table(
        doc,
        [
            ("Get-ChildItem C:\\inetpub\\wwwroot-backups | Sort-Object LastWriteTime -Descending", "Yedek klasörleri tarih sırasıyla gösterir.", "Geri dönüşte hangi yedeğin kullanılacağını seçmek için."),
            ("Copy-Item -LiteralPath <YEDEK_KLASOR>\\* -Destination C:\\inetpub\\wwwroot -Recurse -Force", "Seçilen yedeği canlı klasöre geri kopyalar.", "Öncesinde mevcut wwwroot'u ayrıca yedeklemek gerekir."),
            ("iisreset", "IIS'i yeniden başlatır.", "Rollback sonrası cache/proxy durumunu temizlemek için."),
        ],
        "Rollback Komutları",
    )

    doc.add_heading("15. Terimler Sözlüğü", level=1)
    glossary = [
        ("IIS", "Windows üzerinde web sitesi yayınlayan servis. Bu projede frontend'i yayınlar."),
        ("Next.js static export", "Next.js uygulamasını Node server olmadan IIS'in verebileceği HTML/CSS/JS dosyalarına dönüştürme yöntemi."),
        ("Node.js API", "Login, kullanıcı, iş emri, takvim, rapor ve audit işlemlerini yapan backend."),
        ("PostgreSQL", "Verilerin kalıcı tutulduğu ilişkisel veritabanı."),
        ("Prisma", "TypeScript/Node API ile PostgreSQL arasında ORM katmanı."),
        ("JWT", "Login sonrası API isteklerinde kullanılan token."),
        ("NSSM", "Node uygulamasını Windows Service olarak çalıştırmak için kullanılan araç."),
        ("ARR", "IIS'in reverse proxy yapmasını sağlayan Application Request Routing modülü."),
        ("URL Rewrite", "IIS'te /api gibi gelen yolları başka hedeflere yönlendirmeyi sağlar."),
    ]
    for term, desc in glossary:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{term}: ")
        run.bold = True
        run.font.color.rgb = RGBColor(31, 78, 121)
        p.add_run(desc)

    doc.add_heading("16. Son Kontrol Checklist", level=1)
    for item in [
        "http://127.0.0.1/ 200 dönüyor.",
        "http://127.0.0.1/api/health ok:true ve db:up dönüyor.",
        "HotelOpsApi servisi Running.",
        "postgresql-x64-18 servisi Running.",
        "admin kullanıcısı kurulum parolasıyla dashboard açıyor.",
        "api-service.err boş veya açıklanmış eski hata dışında yeni hata içermiyor.",
        "C:\\inetpub\\wwwroot-backups altında son yayın öncesi yedek var.",
    ]:
        add_bullet(doc, item)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
